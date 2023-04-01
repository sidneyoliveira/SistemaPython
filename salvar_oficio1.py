import os
import docx
import requests
from bs4 import BeautifulSoup
from docx import Document
from pdfminer.high_level import extract_text
from tkinter import filedialog
import customtkinter
from tkinter import *
global pdf_filename
import ctypes

id_item = 0
cnpj_list = ["", "", ""]
razao_list = ["", "", ""]
data_list =  ["", "", ""]



ICON_ERROR = 0x10
ICON_WARNING = 0x30
ICON_QUESTION = 0x20
ICON_INFORMATION = 0x40

def on_input_change(var_name, input_obj):
    def on_change(event):
            globals()[var_name] = input_obj.get()

    return on_change


# Definindo a função que será executada quando o botão for clicado
def input_pdf(var_name, tabela):
    def entry(event):
        global pdf_filename
        pdf_filename = filedialog.askopenfilename(title="Selecione o arquivo MAPA PDF",
                                                  filetypes=(("pdf files", "*.pdf"),))
        salvar_arquivo(pdf_filename, tabela)

    return entry

def exibir_alerta(titulo, mensagem, tipo_icone):
    ctypes.windll.user32.MessageBoxW(0, mensagem, titulo, tipo_icone)

def salvar_arquivo(dados,tabela):

    if dados:
        print(dados)
        # Extrai texto do arquivo PDF PAGINA 1
        pag = extract_text(dados)

        # Divide o texto em linhas
        linhas = pag.split("\n")

        # Imprime todas as linhas
        print(linhas)

        # Imprime a linha de índice 7
        print(linhas[4] + "\n\n")

        # Extrai o número do processo da linha de índice 5
        num = linhas[4].split(" ")[5]
        print(f'numero ofico {num}')

        # Extrai a data do processo da linha de índice 5
        data: str = linhas[10].split(" ")[1]
        print(data)
        dia = data.split('/')[0]
        mes = data.split('/')[1]
        ano = data.split('/')[2]
        print(dia,mes,ano)

        # Extrai o título do processo da linha de índice 7
        titulo = linhas[12].split(" ")
        titulo.remove('DESCRIÇÃO:')
        titulo = " ".join(titulo)
        titulo = titulo.replace('  ', ' ')
        print(titulo)

        # Encontra o índice da linha que contém "Unid. de medida"
        indice_udm = linhas.index("Item Descrição do item")
        # Extrai a descrição até o índice encontrado
        descricao = " ".join(linhas[14:indice_udm])
        descricao = descricao.replace('ESPECIFICAÇÃO/OBJETO:', '')
        descricao = descricao.replace('  ', ' ')
        print(descricao + "\n")

        cnpj_indice = linhas.index("Validade da proposta: 60 dias ")
        print(cnpj_indice)
        print(linhas[cnpj_indice+3])

        cnpj = linhas[cnpj_indice+3].split(": ")[1]
        razao = linhas[cnpj_indice+2].split(": ")[1]

        print(f'cnpj:{cnpj}')
        print(f'razao:{razao}')

        global id_item
        id_item += 1

        cnpj_list[id_item-1] = cnpj
        razao_list[id_item-1] = razao
        data_list[id_item-1] = data





    else:
        # exibir_alerta("Atenção", "Mapa não Importado", 0x30)
        return
    print(cnpj_list)
    print(razao_list)
    print(data_list)
    # Função para extrair dados do cnpj
    numoficio = f'{dia}' + f'{mes}' + '-0001.' + f'{ano}'

    global itens
    itens = ["num", "numoficio", "data", "titulo", "descricao"]
    itens[0] = num
    itens[1] = numoficio
    itens[2] = data
    itens[3] = titulo
    itens[4] = descricao

def salvar_word():

    doc = Document('oficio1.docx')
    tabela = doc.tables[0]
    global itens
    num = itens[0]
    numoficio = itens[1]
    data = itens[2]
    titulo = itens[3]
    descricao = itens[4]

    p = doc.add_paragraph('Oficio nº ')
    p.add_run(numoficio).bold = True
    p.add_run(' - LICITAÇÃO \t\t\t\t\t')

    p.add_run('ITAREMA-CE, ')
    p.add_run(data + '\n\n').bold = True

    doc.add_paragraph('INEZ Helena Braga')
    doc.add_paragraph('Presidente da Comissão de Licitação\n\n')



    p2 = doc.add_paragraph(
            '\t\tConsiderando a realização de pesquisa de preço via e-mail junto ao sistema de cotação pública aCotação processo nº: ')
    p2.add_run(num).bold = True
    p2.add_run(' para: ')
    p2.add_run(descricao).bold = True
    p2.add_run(
            ', encaminha-se ao Setor de Licitação as respectivas propostas juntamente com o mapa de preço médio e comprovações junto ao TCE/CE para providências cabíveis quanto ao seguimento do processo licitatório.\t\t\n')
    p2.alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.JUSTIFY

    table = doc.tables[0]

    referencias = {
        "CNPJ1": f'{cnpj_list[0]}',
        "CNPJ2": f'{cnpj_list[1]}',
        "CNPJ3": f'{cnpj_list[2]}',
        "RAZAO1": f'{razao_list[0]}',
        "RAZAO2": f'{razao_list[1]}',
        "RAZAO3": f'{razao_list[2]}',
        "DATA1": f'{data_list[0]}',
        "DATA2": f'{data_list[1]}',
        "DATA3": f'{data_list[2]}',
        }
    for i in range(len(table.rows)):
        for j in range(len(table.columns)):
            cell = table.cell(i, j)
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    for codigo in referencias:
                        valor = referencias[codigo]
                        new_text = run.text.replace(codigo, valor)
                        run.text = new_text.upper()

    for para in doc.paragraphs:
        for run in para.runs:
            run.font.name = 'Calibri Light'

    tabela = doc.tables[0]

    novo_paragrafo = doc.add_paragraph()

    # Insere a tabela antes do parágrafo vazio
    tabela_antes = novo_paragrafo.insert_paragraph_before('')
    tabela_antes._element.addprevious(tabela._element)

    nome_do_arquivo = f'{numoficio} - {titulo}.docx'
    diretorio = filedialog.askdirectory()
    print(diretorio)
    if diretorio:
        caminho_completo = os.path.join(diretorio, nome_do_arquivo)
        print(caminho_completo)
        doc.save(caminho_completo)
        os.startfile(caminho_completo)



