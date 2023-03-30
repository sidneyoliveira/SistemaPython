import os
import docx
import requests
from bs4 import BeautifulSoup
from docx import Document
from pdfminer.high_level import extract_text
from tkinter import filedialog
import customtkinter
from tkinter import *

def on_input_change(var_name, input_obj):
    def on_change(event):
        globals()[var_name] = input_obj.get()

    return on_change


# Definindo a função que será executada quando o botão for clicado
def input_pdf(var_name):
    def entry(event):
        globals()[var_name] = filedialog.askopenfilename(title="Selecione o arquivo MAPA PDF", filetypes=(("pdf files", "*.pdf"),))
        print(globals()[var_name])
    return entry


def salvar_arquivo(event):
    dados = ''
    print(dados)
    # Extrai texto do arquivo PDF PAGINA 1
    pag1 = extract_text(dados, page_numbers=[0])

    # Divide o texto em linhas
    linhas = pag1.split("\n")

    # Imprime todas as linhas
    print(linhas)

    # Imprime a linha de índice 7
    print(linhas[7] + "\n\n")

   # Extrai o número do processo da linha de índice 5
    num = linhas[5].split(" ")[1]
    print(num)

    # Extrai a data do processo da linha de índice 5
    data: str = linhas[5].split(" ")[4]
    dia = data.split('/')[0]
    mes = data.split('/')[1]
    ano = data.split('/')[2]
    print(data)

    # Extrai o título do processo da linha de índice 7
    titulo = linhas[7].split(" ")
    titulo.remove('DESCRIÇÃO:')
    titulo = " ".join(titulo)
    titulo = titulo.replace('  ', ' ')
    print(titulo)

    # Encontra o índice da linha que contém "Unid. de medida"
    indice_udm = linhas.index("Item")
    # Extrai a descrição até o índice encontrado
    descricao = " ".join(linhas[9:indice_udm])
    descricao = descricao.replace('ESPECIFICAÇÃO:', '')
    descricao = descricao.replace('  ', ' ')
    print(descricao + "\n")

    # Função para extrair dados do cnpj
    def dados_cnpj(a):
        url = f'https://cnpj.biz/{a}'
        response = requests.get(url)
        content = response.content
        dados_cnpj = BeautifulSoup(content, 'html.parser')
        print(dados_cnpj)

        dados_cnpj = dados_cnpj.findAll('b', attrs={'class': 'copy'})
        print(url)
        print(dados_cnpj)
        cnpj = dados_cnpj[0].text
        razao = dados_cnpj[2].text
        return cnpj, razao

    empresa1 = dados_cnpj(cnpj1)
    cnpj1_input = empresa1[0]
    razao1 = empresa1[1]

    empresa2 = dados_cnpj(cnpj2)
    cnpj2_input = empresa2[0]
    razao2 = empresa2[1]

    empresa3 = dados_cnpj(cnpj3)
    cnpj3_input = empresa3[0]
    razao3 = empresa3[1]

    numoficio = f'{dia}' + f'{mes}' + '-0001.' + f'{ano}'

    doc = Document('oficio1.docx')
    tabela = doc.tables[0]

    p = doc.add_paragraph('Oficio nº ')
    p.add_run(numoficio).bold = True
    p.add_run(' - LICITAÇÃO \t\t\t\t\t')

    p.add_run('ITAREMA-CE, ')
    p.add_run(data + '\n\n').bold = True

    doc.add_paragraph('INEZ Helena Braga')
    doc.add_paragraph('Presidente da Comissão de Licitação\n\n')

    p2 = doc.add_paragraph(
        '\t\tConsiderando a realização de pesquisa de preço via e-mail junto ao sistema de cotação pública aCotação processo nº: ')
    p2.add_run(num).bold = True
    p2.add_run(' para: ')
    p2.add_run(descricao).bold = True
    p2.add_run(
        ', encaminha-se ao Setor de Licitação as respectivas propostas juntamente com o mapa de preço médio e comprovações junto ao TCE/CE para providências cabíveis quanto ao seguimento do processo licitatório.\t\t\n')
    p2.alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.JUSTIFY

    table = doc.tables[0]

    referencias = {
        "CNPJ1": f'{cnpj1_input}',
        "CNPJ2": f'{cnpj2_input}',
        "CNPJ3": f'{cnpj3_input}',
        "RAZAO1": f'{razao1}',
        "RAZAO2": f'{razao2}',
        "RAZAO3": f'{razao3}',
        "DATA1": f'{data1}',
        "DATA2": f'{data2}',
        "DATA3": f'{data3}',

    }
    for i in range(len(table.rows)):
        for j in range(len(table.columns)):
            cell = table.cell(i, j)
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    for codigo in referencias:
                        valor = referencias[codigo]
                        new_text = run.text.replace(codigo, valor)
                        run.text = new_text.upper()

    for para in doc.paragraphs:
        for run in para.runs:
            run.font.name = 'Calibri Light'

    tabela = doc.tables[0]

    novo_paragrafo = doc.add_paragraph()

    # Insere a tabela antes do parágrafo vazio
    tabela_antes = novo_paragrafo.insert_paragraph_before('')
    tabela_antes._element.addprevious(tabela._element)

    nome_do_arquivo = f"{numoficio}-{titulo}.docx"
    diretorio = filedialog.askdirectory()
    if diretorio:
        arquivo = docx.Document(nome_do_arquivo)
        caminho_completo = os.path.join(diretorio, nome_do_arquivo)
        arquivo.save(caminho_completo)
        os.startfile(caminho_completo)
