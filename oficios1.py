import os
import docx
import requests
from bs4 import BeautifulSoup
from docx import Document
from pdfminer.high_level import extract_text
from tkinter import filedialog
import customtkinter
from tkinter import *

cnpj1 = ""
cnpj2 = ""
cnpj3 = ""
num = ''
data = ''
numoficio = ''
titulo = ''
descricao = ''

cnpj1_input = ''
cnpj2_input = ''
cnpj3_input = ''
razao1 = ''
razao2 = ''
razao3 = ''
data1 = ''
data2 = ''
data3 = ''
global_var = ''
global doc
doc = ''


class CustomButton(customtkinter.CTkButton):
    def __init__(self, master, **kw):
        customtkinter.CTkButton.__init__(self, master=master, **kw)

    def return_value(self, value):
        self.master.return_value = value


def assign_global_var(value):
    global global_var
    global_var = value


def on_input_change(var_name, input_obj):
    def on_change(event):
        globals()[var_name] = input_obj.get()

    return on_change


# Definindo a função que será executada quando o botão for clicado
def input_pdf():
    try:
        global global_var
        entry = filedialog.askopenfilename(title="Selecione o arquivo MAPA PDF", filetypes=(("pdf files", "*.pdf"),))
        global_var = entry
        print(global_var)
        return global_var
    except FileNotFoundError:
        print("Arquivo não importado.")

def dados_pdf(global_var):
    print(global_var)

    # Extrai texto do arquivo PDF PAGINA 1
    pag1 = extract_text(global_var, page_numbers=[0])

    # Divide o texto em linhas
    linhas = pag1.split("\n")

    # Imprime todas as linhas
    print(linhas)

    # Imprime a linha de índice 7
    print(linhas[7] + "\n\n")
    num = linhas[5].split(" ")[1]
    print(num)

    data: str = linhas[5].split(" ")[4]
    dia = data.split('/')[0]
    mes = data.split('/')[1]
    ano = data.split('/')[2]
    numoficio = f'{dia}' + f'{mes}' + '-0001.' + f'{ano}'
    print(data)

    # Extrai o título do processo da linha de índice 7
    titulo = linhas[7].split(" ")
    titulo.remove('DESCRIÇÃO:')
    titulo = " ".join(titulo)
    titulo = titulo.replace('  ', ' ')
    print(titulo)

    # Encontra o índice da linha que contém "Unid. de medida"

    indice_udm = linhas.index("Item")
    # Extrai a descrição até o índice encontrado

    descricao = " ".join(linhas[9:indice_udm])
    descricao = descricao.replace('ESPECIFICAÇÃO: ', '')
    descricao = descricao.replace('  ', ' ')
    print(descricao + "\n")

    return num, data, numoficio, titulo, descricao


# Função para extrair dados do cnpj
def dados_cnpj(a):
    url = f'https://cnpj.biz/{a}'
    response = requests.get(url)
    content = response.content
    dados_cnpj = BeautifulSoup(content, 'html.parser')
    print(dados_cnpj)

    dados_cnpj = dados_cnpj.findAll('b', attrs={'class': 'copy'})
    print(url)
    print(dados_cnpj)
    cnpj = dados_cnpj[0].text
    razao = dados_cnpj[2].text
    return cnpj, razao


root = customtkinter.CTk()
root.geometry("600x300")
root.title("Oficios")

input_cnpj1 = customtkinter.CTkEntry(root, placeholder_text="CNPJ",
                                     width=200, height=40,
                                     border_width=2,
                                     border_color='#dddddd',
                                     corner_radius=10, )
input_cnpj1.grid(column=1, row=1, padx=10, pady=10)
input_cnpj1.bind("<FocusOut>", on_input_change("cnpj1", input_cnpj1))

input_data1 = customtkinter.CTkEntry(root, placeholder_text="Data",
                                     width=200, height=40,
                                     border_width=2,
                                     border_color='#dddddd',
                                     corner_radius=10, )
input_data1.grid(column=2, row=1, padx=10, pady=10)
input_data1.bind("<FocusOut>", on_input_change("data1", input_data1))

input_cnpj2 = customtkinter.CTkEntry(root, placeholder_text="CNPJ",
                                     width=200, height=40,
                                     border_width=2,
                                     border_color='#dddddd',
                                     corner_radius=10, )

input_cnpj2.grid(column=1, row=2, padx=10, pady=10)
input_cnpj2.bind("<FocusOut>", on_input_change("cnpj2", input_cnpj2))

input_data2 = customtkinter.CTkEntry(root, placeholder_text="Data",
                                     width=200, height=40,
                                     border_width=2,
                                     border_color='#dddddd',
                                     corner_radius=10, )
input_data2.grid(column=2, row=2, padx=10, pady=10)
input_data2.bind("<FocusOut>", on_input_change("data2", input_data2))

input_cnpj3 = customtkinter.CTkEntry(root, placeholder_text="CNPJ",
                                     width=200, height=40,
                                     border_width=2,
                                     border_color='#dddddd',
                                     corner_radius=10, )
input_cnpj3.grid(column=1, row=3, padx=10, pady=10)
input_cnpj3.bind("<FocusOut>", on_input_change("cnpj3", input_cnpj3))

input_data3 = customtkinter.CTkEntry(root, placeholder_text="Data",
                                     width=200, height=40,
                                     border_width=2,
                                     border_color='#dddddd',
                                     corner_radius=10, )
input_data3.grid(column=2, row=3, padx=10, pady=10)
input_data3.bind("<FocusOut>", on_input_change("data3", input_data2))

# Criando a entrada e o botão personalizado
custom_button = CustomButton(root, text="Abrir Mapa PDF", command=lambda: assign_global_var(input_pdf()))
custom_button.grid(column=1, row=4, padx=10, pady=10, sticky=W)

if cnpj1:
    empresa1 = dados_cnpj(cnpj1)
    cnpj1_input = empresa1[0]
    razao1 = empresa1[1]

if cnpj2:
    empresa2 = dados_cnpj(cnpj2)
    cnpj2_input = empresa2[0]
    razao2 = empresa2[1]

if cnpj3:
    empresa3 = dados_cnpj(cnpj3)
    cnpj3_input = empresa3[0]
    razao3 = empresa3[1]

dados_cnpj()

if global_var:
    dados_doc = dados_pdf(global_var)
    print(dados_pdf)
    num = dados_doc[0]
    data = dados_doc[1]
    numoficio = dados_doc[2]
    titulo = dados_doc[3]
    descricao = dados_doc[4]


def criar_documento(dados, dados_cnpj):

    dados = dados_pdf(global_var)
    print(dados)
    numoficio = dados[2]
    print(numoficio)

    titulo = dados[3]
    print(titulo)

    doc = Document('oficio1.docx')
    tabela = doc.tables[0]

    p = doc.add_paragraph('Oficio nº ')
    p.add_run(numoficio).bold = True
    p.add_run(' - LICITAÇÃO \t\t\t\t\t')

    p.add_run('ITAREMA-CE, ')
    p.add_run(data + '\n\n').bold = True

    doc.add_paragraph('INEZ Helena Braga')
    doc.add_paragraph('Presidente da Comissão de Licitação\n\n')

    p2 = doc.add_paragraph(
        '\t\tConsiderando a realização de pesquisa de preço via e-mail junto ao sistema de cotação pública aCotação processo nº: ')
    p2.add_run(num).bold = True
    p2.add_run(' para: ')
    p2.add_run(descricao).bold = True
    p2.add_run(
        ', encaminha-se ao Setor de Licitação as respectivas propostas juntamente com o mapa de preço médio e comprovações junto ao TCE/CE para providências cabíveis quanto ao seguimento do processo licitatório.\t\t\n')
    p2.alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.JUSTIFY

    referencias = {
        "CNPJ1": f'{cnpj1_input}',
        "CNPJ2": f'{cnpj2_input}',
        "CNPJ3": f'{cnpj3_input}',
        "RAZAO1": f'{razao1}',
        "RAZAO2": f'{razao2}',
        "RAZAO3": f'{razao3}',
        "DATA1": f'{data1}',
        "DATA2": f'{data2}',
        "DATA3": f'{data3}',
    }

    for i in range(len(tabela.rows)):
        for j in range(len(tabela.columns)):
            cell = tabela.cell(i, j)
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    for codigo in referencias:
                        valor = referencias[codigo]
                        new_text = run.text.replace(codigo, valor)
                        run.text = new_text.upper()

    for para in doc.paragraphs:
        for run in para.runs:
            run.font.name = 'Calibri Light'

    novo_paragrafo = doc.add_paragraph()

    # Insere a tabela antes do parágrafo vazio
    tabela_antes = novo_paragrafo.insert_paragraph_before('')
    tabela_antes._element.addprevious(tabela._element)

    diretorio = filedialog.askdirectory()
    print(diretorio)

    nome_do_arquivo = f"{numoficio} - {titulo}.docx"
    print(nome_do_arquivo)

    caminho_completo = os.path.join(diretorio, nome_do_arquivo)
    print(caminho_completo)

    doc.save(caminho_completo)
    os.startfile(caminho_completo)


custom_button = customtkinter.CTkButton(root, text="Salvar PDF", command=criar_documento(dados))
custom_button.grid(column=1, row=5, padx=10, pady=10, sticky=W)

# Iniciando a janela principal e esperando pelo retorno de valor
root.mainloop()
